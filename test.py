from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# File path
output_path = "brochure.docx"

# Create a new document
doc = Document()

# Title
title = doc.add_paragraph("🚀 RYX - Digital Solutions & Marketing Guide")
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.runs[0].font.size = Pt(22)
title.runs[0].font.bold = True

subtitle = doc.add_paragraph("Production-Ready Solutions in Days, Not Months")
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(10, 132, 255)

doc.add_paragraph("2024 Professional Business Overview").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_page_break()

# Company Overview
doc.add_heading("COMPANY OVERVIEW", level=1)
doc.add_paragraph("RYX is a premium digital agency specializing in rapid deployment of scalable technology solutions. "
                  "We transform ideas into production-ready applications using cutting-edge technologies and proven methodologies.")

doc.add_paragraph("Mission Statement:", style='Heading 2')
doc.add_paragraph("Democratize access to enterprise-grade technology by delivering production-ready applications and digital experiences at startup speed.")

doc.add_paragraph("Core Values:", style='Heading 2')
doc.add_paragraph("✔ Speed Without Compromise – Deploy in days while maintaining quality\n"
                  "✔ Technical Excellence – 90+ Lighthouse scores, optimized performance\n"
                  "✔ Transparent Pricing – Fixed deliverables, no hidden costs\n"
                  "✔ Client Success – 92% retention and 4.9/5 satisfaction rating")

doc.add_page_break()

# Services
doc.add_heading("SERVICE PORTFOLIO", level=1)

services = [
    ("Micro SaaS Development", "Complete SaaS products from MVP to scale. Includes authentication, payments, dashboards, APIs, database optimization, and DevOps.", "Starting at ₹1,25,000"),
    ("Database Management", "Supabase & MySQL optimization, scaling, automated backups, migration support, and security audits.", "Starting at ₹25,000"),
    ("Web Development", "Landing pages, product sites, full-stack apps with SEO optimization, modern stack, CMS, analytics.", "Starting at ₹45,000"),
    ("AI & Prompt Engineering", "Custom AI workflows, LLM integration, vector DB setup, performance tuning, prompt optimization.", "Starting at ₹75,000"),
    ("Graphic & Brand Design", "Logos, brand identity, brochures, social media graphics, print & digital assets.", "Starting at ₹15,000"),
    ("UI/UX Design", "Research-driven wireframes, interactive prototypes, mobile/web interfaces, usability testing.", "Starting at ₹35,000"),
    ("Illustration & Creative Art", "Custom illustrations, infographics, mascots, campaign artwork, packaging design.", "Starting at ₹20,000")
]

for title, desc, price in services:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)
    p = doc.add_paragraph(price)
    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    p.runs[0].font.highlight_color = None

doc.add_page_break()

# Business Performance
doc.add_heading("BUSINESS PERFORMANCE", level=1)
doc.add_paragraph("Our proven results speak for themselves:")
doc.add_paragraph("📈 120+ Projects Delivered")
doc.add_paragraph("💡 99.9% Uptime SLA")
doc.add_paragraph("🤝 92% Client Retention")
doc.add_paragraph("⭐ 4.9/5 Avg Rating")
doc.add_paragraph("🚀 Project kickoff in < 7 days")

doc.add_page_break()

# Implementation Roadmap
doc.add_heading("IMPLEMENTATION ROADMAP", level=1)
doc.add_paragraph("Our growth and scaling approach is designed in three clear phases.")

doc.add_heading("Phase 1: Foundation (0-3 months)", level=2)
doc.add_paragraph("• Website optimization and SEO\n• Content marketing strategy\n• Onboarding refinement\n• Case study development")

doc.add_heading("Phase 2: Growth (3-6 months)", level=2)
doc.add_paragraph("• LinkedIn lead generation\n• Agency partnerships\n• Advanced AI/Enterprise services\n• Team expansion planning")

doc.add_heading("Phase 3: Scale (6-12 months)", level=2)
doc.add_paragraph("• Premium service tier launch\n• International market expansion\n• White-label platforms\n• Strategic partnerships & acquisitions")

doc.add_page_break()

# Contact Page
doc.add_heading("CONTACT & NEXT STEPS", level=1)
doc.add_paragraph("🌐 Website: https://ryx.dev\n✉️ Email: hello@ryx.dev\n📞 Phone: +91-98765-43210")
doc.add_paragraph("💡 Free Consultation Includes: Scope analysis, tech recommendations, budget estimation, and risk assessment.")

doc.add_heading("Why Choose RYX?", level=2)
doc.add_paragraph("✔ Proven Track Record\n✔ Rapid Deployment\n✔ Transparent Pricing\n✔ Ongoing Support")

doc.add_paragraph("\nBuilt with ❤️ by RYX Team – Deploy production-ready solutions in days, not months.")

# Save the document
doc.save(output_path)

output_path
